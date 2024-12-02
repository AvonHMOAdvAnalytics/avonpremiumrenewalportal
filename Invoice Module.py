import streamlit as st
import pyodbc
import pandas as pd
import datetime as dt
from datetime import datetime
from PIL import Image as Im
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
import re
from azure.storage.blob import BlobServiceClient
import os

#set the page configuration as shown
# st.set_page_config(page_title= 'Invoice Module',layout='wide', initial_sidebar_state='expanded')

#add a image header to the page
# image = Im.open('RenewalPortal.png')
# st.image(image, use_column_width=True)

#write the queries to pull data from the DB
query10 = 'select distinct a.PolicyNo, b.PolicyName, a.FromDate, a.ToDate, a.ClassName\
        from tblClassMaster a\
        left join tbl_Enrollee_Premium b on a.PolicyNo = b.PolicyNo'
query11 = 'select * from tbl_renewal_portal_invoice_module_client_data a\
                where invoiceno = (select max(invoiceno) from tbl_renewal_portal_invoice_module_client_data b where a.Client = b.Client)'

query12 = 'select * from tbl_renewal_portal_invoice_module_plan_data a\
        where invoiceno = (select max(invoiceno) from tbl_renewal_portal_invoice_module_plan_data b where a.Client = b.Client)'

# #define the connection for the DBs
# conn = pyodbc.connect(
#         'DRIVER={ODBC Driver 17 for SQL Server};SERVER='
#         +st.secrets['server']
#         +';DATABASE='
#         +st.secrets['database']
#         +';UID='
#         +st.secrets['username']
#         +';PWD='
#         +st.secrets['password']
#         ) 

# assign the DB credentials to variables
server = os.environ.get('server_name')
database = os.environ.get('db_name')
username = os.environ.get('db_username')
password = os.environ.get('db_password')

# define the DB connection
conn = pyodbc.connect(
        'DRIVER={ODBC Driver 17 for SQL Server};SERVER='
        + server
        +';DATABASE='
        + database
        +';UID='
        + username
        +';PWD='
        + password
        )

#write a function to read the data from the DBs
@st.cache_data(ttl = dt.timedelta(hours=24))
def get_data_from_sql():  
    active_clients = pd.read_sql(query10, conn)
    return active_clients

#assign the data to variables as below
active_clients = get_data_from_sql()

name = st.session_state.get('name', None)
email = st.session_state.get('email', None)


# st.write(active_clients.head())

client = st.sidebar.selectbox(label='Select Client', placeholder='Pick a Client', index=None, options=active_clients['PolicyName'].unique())

unique_plan = active_clients.loc[active_clients['PolicyName'] == client, 'ClassName'].unique()

select_plan = st.sidebar.multiselect('Select Active Plans', unique_plan)

payment_plan = st.sidebar.selectbox(label='Select Client Payment Frequency', options=['Annual', 'Bi-Annual', 'Tri-Annual', 'Quarterly'])
if payment_plan == 'Quarterly':
    payment_period = st.sidebar.selectbox(label='Select Payment Period', options=['1st Quarter', '2nd Quarter', '3rd Quarter', '4th Quarter'])
elif payment_plan == 'Bi-Annual':
    payment_period = st.sidebar.selectbox(label='Select Payment Period', options=['1st Half', '2nd Half'])
elif payment_plan == 'Tri-Annual':
    payment_period = st.sidebar.selectbox(label='Select Payment Period', options=['1st Third', '2nd Third', '3rd Third'])
else:
    payment_period = 'Annual'

#function to generate the period based on the payment plan
def generate_periods(start, end, period):
        # Calculate the start date for the next year's policy period
        next_year_start = pd.Timestamp(year=start.year + 1, month=start.month, day=start.day)
        next_year_end = pd.Timestamp(year=end.year + 1, month=end.month, day=end.day)
        #return the next year end date in the format 'Month Year'
        next_year_end = next_year_end.strftime('%B %Y')

        if payment_plan == 'Quarterly':
                periods = pd.date_range(start = next_year_start, periods=4, freq='3M').strftime('%B %Y')
                if payment_period == '1st Quarter': 
                    period = f'{periods[0]} - {periods[1]}'
                elif payment_period == '2nd Quarter':
                    period = f'{periods[1]} - {periods[2]}'
                elif payment_period == '3rd Quarter':
                    period = f'{periods[2]} - {periods[3]}'
                elif payment_period == '4th Quarter':
                    period = f'{periods[3]} - {next_year_end}'
        elif payment_plan == 'Bi-Annual':
                periods = pd.date_range(start = next_year_start, periods=2, freq='6M').strftime('%B %Y')
                if payment_period == '1st Half':
                    period = f'{periods[0]} - {periods[1]}'
                elif payment_period == '2nd Half':
                    period = f'{periods[1]} - {next_year_end}'
        elif payment_plan == 'Tri-Annual':
                periods = pd.date_range(start = next_year_start, periods=3, freq='4M').strftime('%B %Y')
                if payment_period == '1st Third':
                    period = f'{periods[0]} - {periods[1]}'
                elif payment_period == '2nd Third':
                    period = f'{periods[1]} - {periods[2]}'
                elif payment_period == '3rd Third':
                    period = f'{periods[2]} - {next_year_end}'
        elif payment_plan == 'Annual':
                # Generate an annual period for the next year
                next_year_end = pd.Timestamp(year=end.year + 1, month=end.month, day=end.day)
                start_month_year = next_year_start.strftime('%B %Y')
                end_month_year = next_year_end.strftime('%B %Y')
                period = f'{start_month_year} - {end_month_year}'
        else:
                raise ValueError("Invalid payment plan provided.")
        return period

# Form to input quantity, price, and category for each selected plan
def generate_input_fields(client, plan):
        plan_renewal = []
        for i, plan in enumerate(select_plan):
                i_num_lives_key = f'{plan}_ilives'
                i_premium_key = f'{plan}_ipremium'
                f_num_lives_key = f'{plan}_flives'
                f_premium_key = f'{plan}_fpremium'
                st.subheader(f"Renewal Details for {plan}")
                i_num_lives = st.number_input(f"Enter Number of Individual Lives on {plan}", key=i_num_lives_key, min_value=0, step=1)
                i_premium = st.number_input(f"Enter Final Negotiated Price Per Individual on {plan}", key=i_premium_key, min_value=0.0, step=0.01)
                f_num_lives = st.number_input(f"Enter Number of Family Lives on {plan}", key=f_num_lives_key, min_value=0, step=1)
                f_premium = st.number_input(f"Enter Final Negotiated Price Per Family on {plan}", key=f_premium_key, min_value=0.0, step=0.01)

                # Generate the period based on the payment plan
                periods = generate_periods(policy_start_date, policy_end_date, payment_plan)

                if i_num_lives is not None and i_num_lives > 0:
                        if payment_plan == 'Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'INDIVIDUAL', 'Count': i_num_lives, 'Annual Rate (#)':i_premium, 'Period': periods, 'InvPeriod': payment_period, 'Annual Amount (#)': i_num_lives*i_premium})
                        elif payment_plan == 'Bi-Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'INDIVIDUAL', 'Count': i_num_lives, 'Bi-Annual Rate (#)':i_premium/2, 'Period': periods, 'InvPeriod': payment_period, 'Bi-Annual Amount (#)': i_num_lives*(i_premium/2)})
                        elif payment_plan == 'Tri-Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'INDIVIDUAL', 'Count': i_num_lives, 'Tri-Annual Rate (#)':i_premium/3, 'Period': periods, 'InvPeriod': payment_period, 'Tri-Annual Amount (#)': i_num_lives*(i_premium/3)})   
                        elif payment_plan == 'Quarterly':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'INDIVIDUAL', 'Count': i_num_lives, 'Quarterly Rate (#)':i_premium/4, 'Period': periods, 'InvPeriod': payment_period, 'Quarterly Amount (#)': i_num_lives*(i_premium/4)})
                if f_num_lives is not None and f_num_lives > 0:
                        if payment_plan == 'Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'FAMILY', 'Count': f_num_lives, 'Annual Rate (#)':f_premium, 'Period': periods, 'InvPeriod': payment_period, 'Annual Amount (#)': f_num_lives*f_premium})
                        elif payment_plan == 'Bi-Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'FAMILY', 'Count': f_num_lives, 'Bi-Annual Rate (#)':f_premium/2, 'Period': periods, 'InvPeriod': payment_period, 'Bi-Annual Amount (#)': f_num_lives*(f_premium/2)})
                        elif payment_plan == 'Tri-Annual':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'FAMILY', 'Count': f_num_lives, 'Tri-Annual Rate (#)':f_premium/3, 'Period': periods, 'InvPeriod': payment_period, 'Tri-Annual Amount (#)': f_num_lives*(f_premium/3)})
                        elif payment_plan == 'Quarterly':
                              plan_renewal.append({'PlanID':policyno + '-' + str(i+1), 'client':client, 'Plan': plan, 'Type': 'FAMILY', 'Count': f_num_lives, 'Quarterly Rate (#)':f_premium/4, 'Period': periods, 'InvPeriod': payment_period, 'Quarterly Amount (#)': f_num_lives*(f_premium/4)})
              
                                           
        return plan_renewal

#function to generate the invoice number
def generate_invoice_number():
        invoiced_clients = pd.read_sql(query11, conn)
        policy_invoices = invoiced_clients.loc[invoiced_clients['PolicyNo'] == policyno]
        if not policy_invoices.empty:
               #extract the last sequence of digits from the 'invoiceno' col
               policy_invoices['LastDigits'] = policy_invoices['invoiceno'].str.extract(r'(\d+$)').astype(int)
               #sort the df by the extracted digits to get the highest
               policy_invoices = policy_invoices.sort_values(by='LastDigits',ascending=False)
               #get the most recent invoice id
               invoice_id = policy_invoices['invoiceno'].iloc[0]
               #extract the last digits of the most recent id
               last_digits = re.search(r'\d+$', invoice_id)
               #increment the last digits for the new invoice by 1 
               if last_digits:
                       incremented_value = int(last_digits.group(0)) + 1
                       #replace the last digits with the increment value
                       new_invoice_number =re.sub(r'\d+$', str(incremented_value).zfill(2), invoice_id)
               else:
                       new_invoice_number = f'AVON/{client_region}/{current_day}/{current_month}/00' 
        else:
                new_invoice_number = f'AVON/{client_region}/{current_day}/{current_month}/00'

        return new_invoice_number    


def add_thousand_separators(value):
    """
    Add thousand separators to numeric values.
    """
    try:
        value = float(value)
        return '{:,.2f}'.format(value)  # Format as float with thousand separators and two decimal places
    except ValueError:
        return value  # Return non-numeric values as is


def generate_invoice(inv_data, table_data, payment_plan):
    # Load the DOCX template
    doc = Document("Client_Renewal_Invoice_Template.docx")
    
    # Replace non-table placeholders
    for p in doc.paragraphs:
        if "{InvDate}" in p.text:
            p.text = p.text.replace("{InvDate}", inv_data['InvDate'])
        if "{InvNumber}" in p.text:
            p.text = p.text.replace("{InvNumber}", inv_data['InvNumber'])
        if "{InvTitle}" in p.text:
            p.text = p.text.replace("{InvTitle}", inv_data['InvTitle'])
        if "{InvClient}" in p.text:
            p.text = p.text.replace("{InvClient}", inv_data['InvClient'])
        if "{InvAddress}" in p.text:
            p.text = p.text.replace("{InvAddress}", inv_data['InvAddress'])
        if "{InvState}" in p.text:
            p.text = p.text.replace("{InvState}", inv_data['InvState'])
        if "{InvStartYear}" in p.text:
            p.text = p.text.replace("{InvStartYear}", str(inv_data['InvStartYear']))
            # Increase the font size of the InvQuarter to 14 and make it bold
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        if "{InvEndYear}" in p.text:
            p.text = p.text.replace("{InvEndYear}", str(inv_data['InvEndYear']))
            # Increase the font size of the InvQuarter to 14 and make it bold
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        if "{InvPaymentPeriod}" in p.text:
            p.text = p.text.replace("{InvPaymentPeriod}", inv_data['InvPaymentPeriod'])
            # Increase the font size of the InvQuarter to 14 and make it bold
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        if "{InvPeriod}" in p.text:
            p.text = p.text.replace("{InvPeriod}", inv_data['InvPeriod'])
            # Increase the font size of the InvQuarter to 14 and make it bold
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        
        

    # Set the dynamic table headers based on the payment plan
    if payment_plan == 'Annual':
        headers = ['Plan', 'Type', 'Count', 'Annual Rate (#)', 'Annual Amount (#)']
    elif payment_plan == 'Bi-Annual':
        headers = ['Plan', 'Type', 'Count', 'Bi-Annual Rate (#)', 'Bi-Annual Amount (#)']
    elif payment_plan == 'Tri-Annual':
        headers = ['Plan', 'Type', 'Count', 'Tri-Annual Rate (#)', 'Tri-Annual Amount (#)']
    elif payment_plan == 'Quarterly':
        headers = ['Plan', 'Type', 'Count', 'Quarterly Rate (#)', 'Quarterly Amount (#)']
    else:
        raise ValueError("Invalid payment plan provided.")
    
    # Locate the table in the document
    table = doc.tables[0]  # Assuming there's only one table, or find it by content
    
    # Set the table headers
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # Populate the table with the data
    for i, row in table_data.iterrows():
        new_row = table.add_row().cells
        for j, header in enumerate(headers):
            cell = new_row[j]
            cell.text = "{:,.2f}".format(row[header]) if 'Rate' in header or 'Amount' in header else str(row[header])
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
   
    # Populate the second table
    table2 = doc.tables[1]
    table2.cell(0, 0).text = f'{generate_periods(policy_start_date, policy_end_date, payment_plan)} Premium Invoice'
    table2.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table2.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Compute totals
    total_amount = table_data[headers[-1]].sum()  # Assuming the last header is the Amount column
    nhis_fee = total_amount * 0.01
    grand_total = total_amount + nhis_fee

    # Format totals
    total_amount = "{:,.2f}".format(total_amount)
    nhis_fee = "{:,.2f}".format(nhis_fee)
    grand_total = "{:,.2f}".format(grand_total)
    
    # Populate totals in the second table
    table2.cell(0, 2).paragraphs[0].add_run(total_amount).bold = True
    table2.cell(1, 2).paragraphs[0].add_run(nhis_fee).bold = False
    table2.cell(2, 2).paragraphs[0].add_run(grand_total).bold = True

    # Save the modified document to a BytesIO object and convert to a PDF file
    output = BytesIO()
    doc.save(output)
    output.seek(0)
     
    return output

if client is not None and select_plan is not None:
        #get the policy start and end dates
        policy_start_date = pd.to_datetime(active_clients.loc[active_clients['PolicyName'] == client, 'FromDate'].values[0])
        policy_start_year = policy_start_date.year
        policy_end_date = pd.to_datetime(active_clients.loc[active_clients['PolicyName'] == client, 'ToDate'].values[0])
        policy_end_year = policy_end_date.year
        policyno = str(active_clients.loc[active_clients['PolicyName'] == client, 'PolicyNo'].values[0])

        #use st.form to generate the input fields
        with st.form(key='renewal_form'):
                renewal_data = generate_input_fields(client,select_plan)
                renewal_df = pd.DataFrame(renewal_data)
                displayed_renewal_df = renewal_df.iloc[:, 2:]
                #display the form to input the contact person details
                st.subheader(f"Fill {client}'s Contact Person Details Below")
                client_region = st.selectbox(label='Select Client Region', options=['Lagos', 'South-West', 'South-East', 'North', 'South-South'])
                contact_person_title = st.text_input(label="Title of the Contact Person", help='e.g. CEO, HR Head')
                line_address = st.text_input(label="Street No and Street Name of the Client's Address")
                city = st.text_input(label='Town or City the Client is Located')
                state = st.text_input(label='State the Client is Located')
                recipient_email = st.text_input(label="Contact Person's Email")


                #add a submit button to submit the form
                submitted = st.form_submit_button(label='Preview')
                if submitted:
                        generate_periods(policy_start_date, policy_end_date, payment_plan)
                        #calculate the total premium, nhis fee and grand total
                        num_active_plans = renewal_df['Plan'].nunique()
                        total_lives = renewal_df['Count'].sum()
                        total_premium = renewal_df.iloc[:, -1].sum()
                        nhis_fee = 0.01*total_premium
                        grand_total = total_premium + nhis_fee

                        summary_df = pd.DataFrame({
                                'Item': ['Total Premium', '1% NHIS Fee', 'Grand Total'],
                                'Amount': [total_premium, nhis_fee, grand_total]
                                                })
                        
                        current_date = datetime.now()
                        current_month = current_date.month
                        current_day = current_date.day
                        formatted_date = current_date.strftime('%B %d, %Y')
                        #display the renewal data
                        # displayed_renewal_df = renewal_df.iloc[:, 2:]
                        st.write(displayed_renewal_df)
                        st.write(summary_df)

        # generate_invoice = st.button(label=f'Generate {client} invoice as PDF File')
        # if generate_invoice:
                #calculate the total premium, nhis fee and grand total
        num_active_plans = renewal_df['Plan'].nunique()
        total_lives = renewal_df['Count'].sum()
        total_premium = renewal_df.iloc[:, -1].sum()
        nhis_fee = 0.01*total_premium
        grand_total = total_premium + nhis_fee

        summary_df = pd.DataFrame({
                'Item': ['Total Premium', '1% NHIS Fee', 'Grand Total'],
                'Amount': [total_premium, nhis_fee, grand_total]
                                })
        current_date = datetime.now()
        current_month = current_date.month
        current_day = current_date.day
        formatted_date = current_date.strftime('%B %d, %Y')

        inv_data = {
                'InvDate': formatted_date,
                'InvNumber': generate_invoice_number(),
                'InvTitle': contact_person_title,
                'InvClient': client,
                'InvAddress': line_address + ' ' + city,
                'InvState': state,
                'InvPaymentPeriod': payment_period,
                'InvStartYear': policy_start_year,
                'InvEndYear': policy_end_year,
                'InvPeriod': generate_periods(policy_start_date, policy_end_date, payment_plan)
        }
        
        # Generate the invoice
        invoice = generate_invoice(inv_data, renewal_df, payment_plan)

        #convert the invoice to a PDF file before downloading
        # pdf_invoice = BytesIO()
        # invoice.save(pdf_invoice)
        # pdf_invoice.seek(0)
        # pdf_invoice = pdf_invoice.getvalue()
        # pdf_invoice = BytesIO(pdf_invoice)
        # pdf_invoice = st.download_button(label='Generate Invoice', data=pdf_invoice, file_name=f"{client}_Invoice.pdf", mime='application/pdf')
        
        generate_inv = st.download_button(label='Generate Invoice', data=invoice, file_name=f"{client}_Invoice.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # write the data to the DB
        if generate_inv:
                cursor = conn.cursor()
                try:
                        for plan in renewal_data:
                                cursor.execute('INSERT INTO tbl_renewal_portal_invoice_module_plan_data\
                                                (PlanID, Client, PlanName, PremiumType, Count, PremiumAmt, PaymentPlan, Period, invoiceno, rate, TotalPremium, InvPeriod)\
                                                VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                                                (plan['PlanID'],
                                                plan['client'],
                                                plan['Plan'],
                                                plan['Type'],
                                                plan['Count'],
                                                plan['Annual Rate (#)'] if 'Annual Rate (#)' in plan else (plan['Bi-Annual Rate (#)'] * 2 if 'Bi-Annual Rate (#)' in plan else plan['Tri-Annual Rate (#)'] * 3 if 'Tri-Annual Rate (#)' in plan else plan['Quarterly Rate (#)'] * 4),
                                                'Annual' if 'Annual Rate (#)' in plan else ('Bi-Annual' if 'Bi-Annual Rate (#)' in plan else 'Tri-Annual' if 'Tri-Annual Rate (#)' in plan else 'Quarterly'),
                                                plan['Period'],
                                                generate_invoice_number(),
                                                plan['Annual Rate (#)'] if 'Annual Rate (#)' in plan else (plan['Bi-Annual Rate (#)'] / 2 if 'Bi-Annual Rate (#)' in plan else plan['Tri-Annual Rate (#)'] / 3 if 'Tri-Annual Rate (#)' in plan else plan['Quarterly Rate (#)'] / 4),
                                                plan['Annual Amount (#)'] if 'Annual Amount (#)' in plan else (plan['Bi-Annual Amount (#)'] if 'Bi-Annual Amount (#)' in plan else plan['Tri-Annual Amount (#)'] if 'Tri-Annual Amount (#)' in plan else plan['Quarterly Amount (#)']),
                                                plan['InvPeriod']
                                                )
                                )                
                        cursor.execute('INSERT INTO tbl_renewal_portal_invoice_module_client_data\
                                        (invoiceno, policyno,client,active_plans,total_lives,total_premium,nhis_fee,grand_total,contact_person_title,address,state,recipient_email,submitted_date,clientmgr,clientmgremail,PaymentPeriod,PaymentPlan,InvPeriod)\
                                        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
                                        (generate_invoice_number(),
                                        policyno,
                                        client,
                                        int(num_active_plans),
                                        int(total_lives),
                                        total_premium,
                                        nhis_fee,
                                        grand_total,
                                        contact_person_title,
                                        line_address + ' ' + city,
                                        state,
                                        recipient_email,
                                        datetime.now(),
                                        name,
                                        email,
                                        generate_periods(policy_start_date, policy_end_date, payment_plan),
                                        payment_plan,
                                        payment_period
                                        )
                        )

                        #commit to insert the data into respective tables
                        conn.commit()
                        #display the text after the successful writing of the data to the DB.
                        st.success(f'{client} Invoiced Data Successfully written to the DB, Check your Download Folder for the Invoice')
                        #save the generated invoice to a container on azure blob storage
                        # conn_str = st.secrets['azure_conn_str']
                        conn_str = os.environ.get('azure_conn_str')
                        container_name = 'client-renewal-portal-invoices'
                        blob_service_client = BlobServiceClient.from_connection_string(conn_str)
                        #create a folder in the container with the invoice year containing a subfolder with the invoice month, client manager and client name
                        folder_name = f'{current_date.year}/{current_date.strftime("%B")}/{name}/{client}'
                        blob_client = blob_service_client.get_blob_client(container=container_name, blob=folder_name)
                        blob_client.upload_blob(invoice.getvalue())

                except Exception as e:
                                st.error(f"An error occurred: {str(e)}")
                                conn.rollback()
else:
        #display a warning message if the client and plan are not selected
        st.warning('Please select a client and all active plan(s) to proceed')





                             

