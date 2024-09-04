import streamlit as st
import pyodbc
import pandas as pd
import datetime as dt
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
import os

#set the page configuration
# st.set_page_config(page_title= 'Premium Calculator',layout='wide', initial_sidebar_state='expanded')

#add a image header to the page
image = Image.open('RenewalPortal.png')
st.image(image, use_column_width=True)

#write the queries to pull data from the DB
query12 = 'select MemberNo, PolicyNo, PolicyName, PolicyStartDate, PolicyEndDate, EnrolleeName, PlanType, PremiumType, Enrollmentdate, StopDeleteDate,\
        DATEDIFF(month, PolicyStartDate, StopDeleteDate) [PolicyAge@Deletion],\
        DATEDIFF(month, Enrollmentdate, StopDeleteDate) [PeriodSpentonPolicy]\
        from tbl_EnroleeMember_stg\
        where convert(date,PolicyEndDate) >= convert(date, getdate())\
        and StopDeleteDate > PolicyStartDate\
        and DATEDIFF(month, PolicyStartDate, StopDeleteDate) != DATEDIFF(month, Enrollmentdate, StopDeleteDate)'

query13 = 'select distinct a.PolicyNo, b.PolicyName, a.FromDate, a.ToDate, a.ClassName\
        from tblClassMaster a\
        join tblEnrolleePremium b on a.PolicyNo = b.PolicyNo\
        where convert(date,a.ToDate) >= convert(date, getdate())'

query14 = "select MemberNo, MemberHeadNo, PolicyNo, PolicyName, PolicyStartDate, PolicyEndDate, EnrolleeName, PlanType, PremiumType, Enrollmentdate, StopDeleteDate,\
        DATEDIFF(month, PolicyStartDate, EnrollmentDate) [PolicyAge@Enrollment],\
        DATEDIFF(month, Enrollmentdate, PolicyEndDate) +1 [MonthsLeftonPolicy]\
        from tbl_EnroleeMember_stg\
        where convert(date,PolicyEndDate) >= convert(date, getdate())\
        and EnrollmentDate > PolicyStartDate\
        and StopDeleteDate = '1900-01-01 00:00:00.000'"

query15 = 'select * from tbl_renewal_portal_invoice_module_plan_data a\
    where invoiceno = (select max(invoiceno) from tbl_renewal_portal_invoice_module_plan_data b where a.Client = b.Client)'

query16 = 'select * from tbl_renewal_portal_invoice_module_client_data a\
        where invoiceno = (select max(invoiceno) from tbl_renewal_portal_invoice_module_client_data b where a.Client = b.Client)'

# #define the connection for the DBs
# conn = pyodbc.connect(
#         'DRIVER={ODBC Driver 17 for SQL Server};SERVER='
#         +st.secrets['server']
#         +';DATABASE='
#         +st.secrets['database']
#         +';UID='
#         +st.secrets['username']
#         +';PWD='
#         +st.secrets['password']
#         ) 

# assign the DB credentials to variables
server = os.environ.get('server_name')
database = os.environ.get('db_name')
username = os.environ.get('db_username')
password = os.environ.get('db_password')

# define the DB connection
conn = pyodbc.connect(
        'DRIVER={ODBC Driver 17 for SQL Server};SERVER='
        + server
        +';DATABASE='
        + database
        +';UID='
        + username
        +';PWD='
        + password
        )

#write a function to read the data from the DBs
@st.cache_data(ttl = dt.timedelta(hours=24))
def data_loading():  
    active_clients = pd.read_sql(query13, conn)
    deleted_enrollees = pd.read_sql(query12, conn)
    added_enrollees = pd.read_sql(query14, conn)
    return active_clients, deleted_enrollees, added_enrollees

def generate_invoice(inv_data, table_data):
    # Load the DOCX template
    doc = Document("Recon_Invoice_Template.docx")
    
    # Replace non-table placeholders
    for p in doc.paragraphs:
        if "{InvDate}" in p.text:
            p.text = p.text.replace("{InvDate}", inv_data['InvDate'])
        if "{InvNumber}" in p.text:
            p.text = p.text.replace("{InvNumber}", inv_data['InvNumber'])
        if "{InvTitle}" in p.text:
            p.text = p.text.replace("{InvTitle}", inv_data['InvTitle'])
        if "{InvClient}" in p.text:
            p.text = p.text.replace("{InvClient}", inv_data['InvClient'])
        if "{InvAddress}" in p.text:
            p.text = p.text.replace("{InvAddress}", inv_data['InvAddress'])
        if "{InvState}" in p.text:
            p.text = p.text.replace("{InvState}", inv_data['InvState'])
        if "{InvQuarter}" in p.text:
            p.text = p.text.replace("{InvQuarter}", inv_data['InvQuarter'])
            #increase the font size of the InvQuarter to 14 and make it bold
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        
    
    #aggregate the dat for the second table
    total_amount = table_data['TotalPremium'].sum()
    nhis_fee = total_amount * 0.01
    grand_total = total_amount + nhis_fee

    #format the total amount, nhis fee and grand total to 2 decimal places and add a comma separator
    total_amount = "{:,.2f}".format(total_amount)
    nhis_fee = "{:,.2f}".format(nhis_fee)
    grand_total = "{:,.2f}".format(grand_total)

    # Locate the table in the document
    table = doc.tables[0]  # Assuming there's only one table, or find it by content
    
    # Remove the template row(s)
    for row in table.rows[1:]:
        tbl = table._tbl
        tbl.remove(row._tr)
    
    # Insert dynamic rows
    for index, row in table_data.iterrows():
        row_cells = table.add_row().cells
        #format the total premium to 2 decimal places and add a comma separator
        row['TotalPremium'] = "{:,.2f}".format(row['TotalPremium'])
        row['ProratePremium'] = "{:,.2f}".format(row['ProratePremium'])
        for i, key in enumerate(['PlanType', 'PremiumType', 'MonthsLeftonPolicy', 'No. of Members Remaining', 'ProratePremium', 'Period', 'TotalPremium']):
            cell = row_cells[i]
            cell.text = str(row.get(key, ''))
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #populate the second table
    table2 = doc.tables[1]
    table2.cell(0, 0).text = f'{recon_period} Reconciliation Invoice'
    table2.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table2.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table2.cell(0, 2).paragraphs[0].add_run(total_amount).bold = True
    table2.cell(1, 2).paragraphs[0].add_run(nhis_fee).bold = False
    table2.cell(2, 2).paragraphs[0].add_run(grand_total).bold = True

    # Save the modified document to a BytesIO object and convert to a pdf file
    output = BytesIO()
    doc.save(output)
    output.seek(0)
     
    return output

def generate_invoice_no():
    client_invoiceno = invoiced_clients.loc[invoiced_clients['client'] == client, 'invoiceno'].values[0]
    if recon_period == 'First Quarter':
        inv_number = f'{client_invoiceno}Q1/Recon'
    elif recon_period == 'Second Quarter':
        inv_number = f'{client_invoiceno}Q2/Recon'
    elif recon_period == 'Third Quarter':
        inv_number = f'{client_invoiceno}Q3/Recon'
    elif recon_period == 'Fourth Quarter':
        inv_number = f'{client_invoiceno}Q4/Recon'
    return inv_number
#assign the data to variables as below
active_clients, deleted_enrollees, added_enrollees = data_loading()

invoiced_clients = pd.read_sql(query16, conn)
invoiced_clients['PolicyNo'] = invoiced_clients['PolicyNo'].astype(str)

invoiced_plan = pd.read_sql(query15, conn)


active_clients['PolicyNo'] = active_clients['PolicyNo'].astype(str)

cols_to_convert = ['MemberNo', 'PolicyNo']

deleted_enrollees[cols_to_convert] = deleted_enrollees[cols_to_convert].astype(int).astype(str)
added_enrollees[cols_to_convert] = added_enrollees[cols_to_convert].astype(int).astype(str)

date_cols = ['PolicyStartDate', 'PolicyEndDate', 'Enrollmentdate', 'StopDeleteDate']
date_format = '%d/%m/%Y'
for col in date_cols:
    try:
        deleted_enrollees[col] = pd.to_datetime(deleted_enrollees[col], format=date_format, errors='raise').dt.date
        added_enrollees[col] = pd.to_datetime(added_enrollees[col], format=date_format, errors='raise').dt.date
    except ValueError as e:
        st.write(f"Error converting column '{col}': {e}")

name = st.session_state.get('name', None)
email = st.session_state.get('email', None)



#create a select box on the sidebar to allow users select a client from the active client list ans assigne to a variable
client = st.sidebar.selectbox(label='Select Client', placeholder='Pick a Client', index=None, options=active_clients['PolicyName'].unique())

recon_period = st.sidebar.selectbox(label='Reconciliation Period',placeholder='Select Reconciliation Period', index=None, options=['First Quarter', 'Second Quarter', 'Third Quarter', 'Fourth Quarter'])

if client is not None and recon_period is not None:
    policyid = str(active_clients.loc[active_clients['PolicyName'] == client, 'PolicyNo'].values[0])

    policy_start = pd.Timestamp(active_clients.loc[active_clients['PolicyNo']==policyid, 'FromDate'].values[0]).date()
    policy_end = pd.Timestamp(active_clients.loc[active_clients['PolicyNo']==policyid, 'ToDate'].values[0]).date()
    added_enrollees['Enrollmentdate'] = pd.to_datetime(added_enrollees['Enrollmentdate'], errors='coerce')

    st.info(f"{client} Current Policy Cycle: {policy_start} - {policy_end}")

    if client in invoiced_plan['Client'].values:
        payment_plan = invoiced_plan.loc[invoiced_plan['Client'] == client, 'PaymentPlan'].values[0]

        if recon_period == 'First Quarter' and payment_plan == 'Annual':
            recon_start = policy_start
            recon_end = policy_start + pd.DateOffset(months=3)
            recon_start = pd.to_datetime(recon_start)
            recon_end = pd.to_datetime(recon_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')

            period_end = policy_start + pd.DateOffset(months=12)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the end of the policy cycle
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']

        elif recon_period == 'First Quarter' and payment_plan == 'Bi-Annual':
            recon_start = policy_start
            recon_end = policy_start + pd.DateOffset(months=3)
            recon_start = pd.to_datetime(recon_start)
            recon_end = pd.to_datetime(recon_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')

            #create a end of bi-annual policy cycle
            period_end = policy_start + pd.DateOffset(months=5)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of bi-annual policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 6 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'First Quarter' and payment_plan == 'Quarterly':
            recon_start = policy_start
            recon_end = policy_start + pd.DateOffset(months=3)
            recon_start = pd.to_datetime(recon_start)
            recon_end = pd.to_datetime(recon_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of quarterly policy cycle
            period_end = policy_start + pd.DateOffset(months=3)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 3 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Second Quarter' and payment_plan == 'Annual':
            recon_start = policy_start + pd.DateOffset(months=3)
            recon_end = policy_start + pd.DateOffset(months=6)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo','EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + policy_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the end of the policy cycle
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Second Quarter' and payment_plan == 'Bi-Annual':
            recon_start = policy_start + pd.DateOffset(months=3)
            recon_end = policy_start + pd.DateOffset(months=6)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of bi-annual policy cycle
            period_end = policy_start + pd.DateOffset(months=5)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of bi-annual policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 6 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Second Quarter' and payment_plan == 'Quarterly':
            recon_start = policy_start + pd.DateOffset(months=3)
            recon_end = policy_start + pd.DateOffset(months=6)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of quarterly policy cycle
            period_end = policy_start + pd.DateOffset(months=3)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 3 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Third Quarter' and payment_plan == 'Annual':
            recon_start = policy_start + pd.DateOffset(months=6)
            recon_end = policy_start + pd.DateOffset(months=9)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + policy_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the end of the policy cycle
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Third Quarter' and payment_plan == 'Bi-Annual':
            recon_start = policy_start + pd.DateOffset(months=6)
            recon_end = policy_start + pd.DateOffset(months=9)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of bi-annual policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + policy_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Third Quarter' and payment_plan == 'Quarterly':
            recon_start = policy_start + pd.DateOffset(months=6)
            recon_end = policy_start + pd.DateOffset(months=9)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of quarterly policy cycle
            period_end = policy_start + pd.DateOffset(months=9)
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 9 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Fourth Quarter' and payment_plan == 'Annual':
            recon_start = policy_start + pd.DateOffset(months=9)
            recon_end = pd.to_datetime(policy_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + policy_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the end of the policy cycle
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Fourth Quarter' and payment_plan == 'Bi-Annual':
            recon_start = policy_start + pd.DateOffset(months=9)
            recon_end = pd.to_datetime(policy_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of bi-annual policy cycle
            period_end = policy_end
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of bi-annual policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        elif recon_period == 'Fourth Quarter' and payment_plan == 'Quarterly':
            recon_start = policy_start + pd.DateOffset(months=9)
            recon_end = pd.to_datetime(policy_end)
            selected_client_added_df = added_enrollees.loc[
                (added_enrollees['PolicyNo'] == policyid) & 
                (added_enrollees['Enrollmentdate'] >= recon_start) & 
                (added_enrollees['Enrollmentdate'] < recon_end),
                ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
            ].set_index('MemberNo')

            selected_client_deleted_df = deleted_enrollees.loc[
                (deleted_enrollees['PolicyNo'] == policyid) & 
                (deleted_enrollees['StopDeleteDate'] >= recon_start) &
                (deleted_enrollees['StopDeleteDate'] < recon_end),
                ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
            ].set_index('MemberNo')
            #create a end of quarterly policy cycle
            period_end = policy_end
            #create a new column in selected_client_added_df to display a duration from enrollment date to the end of the policy cycle e.g. March 2021 - August 2021
            selected_client_added_df['Period'] = selected_client_added_df['Enrollmentdate'].dt.strftime('%B %Y') + ' - ' + period_end.strftime('%B %Y')
            #create a new column in selected_client_added_df to display the number of months from enrollment date to the recon_end
            selected_client_added_df['MonthsLeftonPolicy'] = 12 - selected_client_added_df['PolicyAge@Enrollment']
        total_added_enrollees = selected_client_added_df.shape[0]
        total_deleted_enrollees = selected_client_deleted_df.shape[0]
        num_active_plans = invoiced_plan.loc[invoiced_plan['Client'] == client, 'PlanName'].nunique()
    
    
    else:
        st.warning(f"{client} has not been invoiced for the current policy cycle")
        selected_client_added_df = added_enrollees.loc[
            added_enrollees['PolicyNo'] == policyid,
              ['MemberNo', 'MemberHeadNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Enrollment']
              ].set_index('MemberNo')
        selected_client_deleted_df = deleted_enrollees.loc[
            deleted_enrollees['PolicyNo'] == policyid,
            ['MemberNo', 'EnrolleeName', 'PlanType', 'PremiumType', 'Enrollmentdate', 'StopDeleteDate', 'PolicyAge@Deletion', 'PeriodSpentonPolicy']
        ].set_index('MemberNo')
        selected_client_added_df['Period'] = None
        selected_client_added_df['MonthsLeftonPolicy'] = None
    if selected_client_added_df.empty:
        st.warning(f"No new members added within {recon_period}")
    
    else:
        invoiced_plan['PremiumType'] = invoiced_plan['PremiumType'].str.upper()
        #check for the unique plan types and premium types in selected_client_added_df and retrieve the premium value for each plan type in invoiced_plan
        plan_premiums = invoiced_plan.loc[invoiced_plan['Client'] == client, ['PlanName', 'PremiumType', 'PremiumAmt']].reset_index(drop=True)
       
        plan_premiums = plan_premiums.drop_duplicates(subset=['PlanName', 'PremiumType'])
        #create a new column in plan_premium to display the prorate premium per month for each plan type
        plan_premiums['ProratePremium'] = plan_premiums['PremiumAmt']/12
        plan_premiums = plan_premiums.set_index(['PlanName', 'PremiumType'])['ProratePremium']
       


        st.subheader(f"{client}'s Member(s) Added Within {recon_period}")
        st.write(selected_client_added_df)

        family_cond_df = selected_client_added_df.loc[
            (selected_client_added_df['PremiumType'] == 'FAMILY') &
            (selected_client_added_df['MemberHeadNo'] == selected_client_added_df.index)
             ]
        individual_cond_df = selected_client_added_df.loc[
            (selected_client_added_df['PremiumType'] == 'INDIVIDUAL')
        ]

        eligible_add = pd.concat([family_cond_df, individual_cond_df])
        #group the selected_client_added_df by PlanType and PolicyAge@Enrollment and display in a dataframe
        summary_added_df = eligible_add.groupby(['PlanType','PremiumType', 'PolicyAge@Enrollment' ,'MonthsLeftonPolicy', 'Period']).size().reset_index(name='No. of Added Enrollees')
        #create a new column to display the prorate premium per month for each plan type and member type i.e. individual or family added
        summary_added_df['ProratePremium'] = summary_added_df.set_index(['PlanType', 'PremiumType']).index.map(plan_premiums)
        

        
        # #create a new column to display the period for each line item in the summary_added_df e.g. March 2021 - August 2021
        # summary_added_df['Period'] = recon_start.strftime('%B %Y') + ' - ' + recon_end.strftime('%B %Y')

        #display the summary in a dataframe
        st.subheader(f'Summary of Members Added Within {recon_period}')
        st.write(summary_added_df)

    #check if the selected_client_deleted_df is not empty
    if not selected_client_deleted_df.empty:
        #check for members with policyage@deletion less than 6 months
        eligible_replace = selected_client_deleted_df.loc[selected_client_deleted_df['PolicyAge@Deletion'] < 6]
        #check if eligible_replace is not empty and display the aggregated data in a dataframe
        if not eligible_replace.empty:
            st.subheader('Summary of Members Eligible for Replacement')
            summary_deleted_df = eligible_replace.groupby(['PlanType', 'PremiumType', 'PolicyAge@Deletion']).size().reset_index(name='Count')
            st.write(summary_deleted_df)
        else:
            st.warning('No member is eligible for replacement')
        
         #join the 'count' column in summary_deleted_df to the summary_added_df where the plan type are equal and where the policy age at enrollment is equal to the policy age at deletion
        reconciled_df = summary_added_df.merge(summary_deleted_df, how='left', left_on=['PlanType', 'PremiumType', 'PolicyAge@Enrollment'], right_on=['PlanType', 'PremiumType', 'PolicyAge@Deletion'])
        #create a new column in summary_added_df to display the number of members replaced
        reconciled_df['No. of Replaced Members'] = reconciled_df['Count']
        #fill the NaN values in the 'No. of Replaced Members' column with 0
        reconciled_df['No. of Replaced Members'] = reconciled_df['No. of Replaced Members'].fillna(0)
        #create a new column in summary_added_df to display the number of members remaining after replacement
        reconciled_df['No. of Members Remaining'] = reconciled_df['No. of Added Enrollees'] - reconciled_df['No. of Replaced Members']

        # Variable to hold the total count of matching deleted members
        total_count = 0

        # Iterate through each row in summary_added_df
        for index, added_row in summary_added_df.iterrows():
            # Extract PlanType and PolicyAge@Enrollment for the current row
            plan_type = added_row['PlanType']
            policy_age_enrollment = added_row['PolicyAge@Enrollment']
            premiumtype = added_row['PremiumType']
            
            # Filter summary_deleted_df to find matching PlanType and where PolicyAge@Deletion is less than PolicyAge@Enrollment
            matching_deleted = summary_deleted_df[
                (summary_deleted_df['PlanType'] == plan_type) &
                (summary_deleted_df['PremiumType'] == premiumtype) &
                (summary_deleted_df['PolicyAge@Deletion'] < policy_age_enrollment)
            ]
            
            # Sum the 'Count' for matching records and add to the total
            count_sum = matching_deleted['Count'].sum()
            total_count += count_sum

            #Subtract the total_count from the 'No. of Members Remaining' column in reconciled_df for each plan type
            reconciled_df.loc[
                (reconciled_df['PlanType'] == plan_type) &
                (reconciled_df['PremiumType'] == premiumtype),
                'No. of Members Remaining'
            ] -= count_sum

            #add the total_count to the 'No. of Replaced Members' column in reconciled_df for each plan type
            reconciled_df.loc[
                (reconciled_df['PlanType'] == plan_type) &
                (reconciled_df['PremiumType'] == premiumtype),
                'No. of Replaced Members'
            ] += count_sum

            #create a new column that displays the total premium by multiplying the 'No. of Members Remaining' column by the 'ProratePremium' column and the 'MonthsLeftonPolicy' column
            reconciled_df['TotalPremium'] = reconciled_df['No. of Members Remaining'] * reconciled_df['ProratePremium'] * reconciled_df['MonthsLeftonPolicy']

        #check for plantypes with No. of Members Remaining greater than 0 and display selected columns in a dataframe
        remaining_members = reconciled_df.loc[
            reconciled_df['No. of Members Remaining'] > 0,
            ['PlanType', 'PremiumType', 'MonthsLeftonPolicy', 'No. of Added Enrollees', 'No. of Replaced Members', 'No. of Members Remaining', 'ProratePremium', 'Period', 'TotalPremium']
            ]
        final_recon = remaining_members[['PlanType', 'PremiumType', 'MonthsLeftonPolicy', 'No. of Added Enrollees', 'No. of Replaced Members', 'No. of Members Remaining', 'ProratePremium', 'Period', 'TotalPremium']].copy()
        

    elif selected_client_deleted_df.empty:
        st.warning(f"No members deleted within {recon_period}")
        final_recon = summary_added_df[['PlanType', 'PremiumType', 'MonthsLeftonPolicy', 'No. of Added Enrollees', 'ProratePremium', 'Period']].copy()
        final_recon['No. of Replaced Members'] = 0
        final_recon['No. of Members Remaining'] = final_recon['No. of Added Enrollees']
        final_recon['TotalPremium'] = final_recon['No. of Members Remaining'] * final_recon['ProratePremium'] * final_recon['MonthsLeftonPolicy']
    
    total_premium = final_recon['TotalPremium'].sum()
    total_invoiced_enrollees = final_recon['No. of Members Remaining'].sum()
    total_replaced_enrollees = final_recon['No. of Replaced Members'].sum()
    
    st.subheader('Final Reconciliation Summary for Additional Invoice')
    st.write(final_recon)

    #Use the Recon_Invoice_Template.docx to generate an additional invoice based on the final reconciliation summary
    invoice_data = {
        'InvClient': client,
        'InvDate': dt.datetime.now().strftime('%d/%m/%Y'),
        'InvNumber': generate_invoice_no(),
        'InvAddress': invoiced_clients.loc[invoiced_clients['client'] == client, 'address'].values[0],
        'InvTitle': invoiced_clients.loc[invoiced_clients['client'] == client, 'contact_person_title'].values[0],
        'InvState': invoiced_clients.loc[invoiced_clients['client'] == client, 'state'].values[0],
        'InvQuarter': recon_period,
    }
    
    invoice = generate_invoice(invoice_data, final_recon)
        
    generate = st.download_button(label='Generate Invoice', data=invoice, file_name=f"{client}_Invoice.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    if generate:
        #write the data to the DB
        cursor = conn.cursor()
        try:
            for index, row in final_recon.iterrows():
                cursor.execute(
                    "INSERT INTO tbl_renewalportal_plan_recon_invoice (invoiceno, PlanName, PremiumType, Prorated_Premium, monthsleftonpolicy,\
                    added_enrollees, replaced_enrollees, invoiced_enrollees, invoice_period, total_premium, submission_date)\
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (invoice_data['InvNumber'], row['PlanType'], row['PremiumType'], row['ProratePremium'], int(row['MonthsLeftonPolicy']),
                    int(row['No. of Added Enrollees']), int(row['No. of Replaced Members']), int(row['No. of Members Remaining']), row['Period'],
                    row['TotalPremium'], dt.datetime.now())
                )
            cursor.execute(
                'INSERT INTO tbl_renewalportal_client_recon_invoice (invoiceno, policyno, client, policystartdate, policyenddate,\
                    num_active_plans, total_added_lives, total_deleted_lives, total_replaced_lives, total_invoiced_lives, total_additional_premium,\
                    recon_period, payment_plan, client_mgr, clientmgr_email, submission_date)\
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                    (invoice_data['InvNumber'], policyid, client, policy_start, policy_end, int(num_active_plans), int(total_added_enrollees),
                    int(total_deleted_enrollees), int(total_replaced_enrollees), int(total_invoiced_enrollees), total_premium, recon_period, payment_plan,
                    name, email, dt.datetime.now())
            )
            conn.commit()
            st.success('Invoice generated successfully')
        except Exception as e:
            st.error(f"An error occurred: {e}")
            conn.rollback()

else:
    st.info("Select a client from the list on the sidebar")